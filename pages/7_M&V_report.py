# -*- coding: utf-8 -*-
"""
Created on Thu Sep  8 15:41:43 2022

@author: Bruno Tabet
"""

import numpy as np
import docx
from copy import deepcopy
import os
import streamlit as st
from helpful_funcs.useful_funcs import *
import json
from os.path import exists
from docx import Document
import matplotlib.pyplot as plt
import plotly_express as px
from streamlit_plotly_events import plotly_events
import plotly.graph_objects as go
from plotly.subplots import make_subplots
from docx.shared import Inches
from units import units
import scipy.stats
from docx.oxml.shared import OxmlElement, qn
from docx.text.paragraph import Paragraph
import matplotlib.mathtext as mathtext
from datetime import datetime
from helpful_funcs.useful_funcs import *

import base64
import os
import json
import pickle
import uuid
import re
import git
from git import Repo

initialization(st)

st.header('M&V Report')

if 'name_sidebar' in st.session_state:
    st.sidebar.title("Project name : " + st.session_state['name_sidebar'])

doc = docx.Document('M&VTemplate.docx')

if st.session_state['M&V'] == 1:
    
    st.write("You can now automatically create the M&V report of a chosen project.")
    
    file_name = 'database.json'
    if exists(file_name):
        f = open(file_name)
        st.session_state['db2'] = json.load(f)
        st.session_state['selected_project_name'] = st.selectbox('Select the project for which you want to create an M&V report:',
                                                            options = [value for value in st.session_state['db2']])
    
        if st.button('Confirm choice !'):
            st.session_state['selected_project'] = st.session_state['db2'][st.session_state['selected_project_name']]
            st.session_state['M&V'] = 1.1
            st.experimental_rerun()


if st.session_state['M&V'] == 1.1:
    
    st.selectbox('Select the project for which you want to create an M&V report:',
                 options = [st.session_state['selected_project_name']], disabled = True)
    
    st.session_state['selected_scope'] = st.selectbox('Select a scope from the project '+ st.session_state['selected_project_name'] + '',
                                                      options = [scope for scope in st.session_state['db2'][st.session_state['selected_project_name']]])
    
    if st.button('Confirm choice !'):
        st.session_state['M&V'] = 1.2
        st.experimental_rerun()
    

if st.session_state['M&V'] == 1.2:
    
    st.selectbox('Select the project for which you want to create an M&V report:',
                 options = [st.session_state['selected_project_name']], disabled = True)
    
    st.session_state['selected_scope'] = st.selectbox('Select a scope from the project '+ st.session_state['selected_project_name'] + '',
                                                      options = [st.session_state['selected_scope']], disabled = True)
    
    utilities = []
    
    for i in range (len(st.session_state['selected_project'][st.session_state['selected_scope']])):
        if st.session_state['selected_project'][st.session_state['selected_scope']][i]['Utility']['name'] not in utilities:
            utilities.append(st.session_state['selected_project'][st.session_state['selected_scope']][i]['Utility']['name'])

    st.session_state['selected_utilities'] = st.multiselect('Select the utilities from the project '+ st.session_state['selected_project_name'] + ' in the selected scope ' + st.session_state['selected_scope']+ '',
                                                            options = utilities)
    
    st.write('For each utility, we will pick the latest version.')
    
    if st.button('CHANGE'):
        st.session_state['M&V'] = 1
        st.experimental_rerun()
    
    st.write('')
    st.write('')
    
    
    if len(st.session_state['selected_utilities']) == 0:
        st.button('CREATE REPORT', key = 46546534541534, disabled = True)
        st.write('**Please choose at least one utility.**')
    
    
    else:
        if st.button('CREATE REPORT', key = 989777777):
            
            utilities_already_selected = []
            st.session_state['selected_projects'] = []
            
            
            for i in range (len(st.session_state['selected_project'][st.session_state['selected_scope']])-1, -1, -1):
                if st.session_state['selected_project'][st.session_state['selected_scope']][i]['Utility']['name'] not in utilities_already_selected:
                    if st.session_state['selected_project'][st.session_state['selected_scope']][i]['Utility']['name'] in st.session_state['selected_utilities'] :
                        utilities_already_selected.append(st.session_state['selected_project'][st.session_state['selected_scope']][i]['Utility']['name'])
                        st.session_state['selected_projects'].append(st.session_state['selected_project'][st.session_state['selected_scope']][i])
    
    
            st.session_state['M&V'] = 1.3
            st.experimental_rerun()


if st.session_state['M&V'] == 1.3:
    
    st.write(os.path.abspath(os.getcwd()))
    st.write(os.path.abspath("Final_Report.docx"))
    # st.stop()
    
    with st.spinner('Creating the M&V report'):
        
        projects = st.session_state['selected_projects']
        
        
        def find_paragraph(string):
            paragraph = None
            for par in doc.paragraphs:
                if string in par.text:
                    paragraph = par
            return paragraph
        
        
        def copy_table_after(table, paragraph):
            tbl, p = table._tbl, paragraph._p
            new_tbl = deepcopy(tbl)
            p.addnext(new_tbl)
            return new_tbl

        def find_table(string):
            new_table = None
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if string in cell.text:
                            new_table = table
            return new_table
        
        def delete_paragraph(string):
            for par in doc.paragraphs:
                if string in par.text:
                    par.text = ''
        
        def delete_column_in_table(table, columns):
            grid = table._tbl.find("w:tblGrid", table._tbl.nsmap)
            for cell in table.column_cells(columns):
                cell._tc.getparent().remove(cell._tc)
            col_elem = grid[columns]
            grid.remove(col_elem)
            
        
        def format_str(n):
            
            if type(n) == str:
                return n
            
            if type(n) == int:
                new = "{:,}".format(n)
                return str(new)
            
            new = n
            
            if abs(n) < 10:
                new = round(n, 3)
            elif abs(n) < 100:
                new = round(n, 2)
            elif abs(n) < 1000:
                new = round(n, 1)
            elif abs(n) >= 1000:
                new = round(n, 1)
                new = "{:,}".format(new)
                return str(new)[:-2]
                
            return str(new)
        
        
        
        def insert_paragraph_after(paragraph, text, style=None):

            new_p = OxmlElement("w:p")
            paragraph._p.addnext(new_p)
            new_para = Paragraph(new_p, paragraph._parent)
            if text:
                new_para.add_run(text)
            if style is not None:
                new_para.style = style
            
            return new_para
                
        
        def add_table_caption(location, text, doc):
        
            paragraph = doc.paragraphs[-1]
        
            drapeau = False
            for para in doc.paragraphs:
                if not drapeau:
                    if type(location) == str:
                        if location in para.text:
                            paragraph = insert_paragraph_after(para, "Table ", style="Caption")
                            drapeau = True
                    else:
                        if location.text == para.text:
                            paragraph = insert_paragraph_after(para, "Table ", style="Caption")
                            drapeau = True
                        
        
            run = paragraph.add_run()
            r = run._r
            fldChar = OxmlElement('w:fldChar')
            fldChar.set(qn('w:fldCharType'), 'begin')
            r.append(fldChar)
            instrText = OxmlElement('w:instrText')
            instrText.text = ' SEQ Table \\* ARABIC'
            r.append(instrText)
            fldChar = OxmlElement('w:fldChar')
            fldChar.set(qn('w:fldCharType'), 'end')
            r.append(fldChar)
        
            paragraph.add_run(f": {text}")
        
        
        def add_figure_caption(location, text, doc):
        
            paragraph = doc.paragraphs[-1]
        
            drapeau = False
            for para in doc.paragraphs:
                if not drapeau:
                    if location in para.text:
                        paragraph = insert_paragraph_after(para, "Figure ", style="Caption")
                        drapeau = True
        
            run = paragraph.add_run()
            r = run._r
            fldChar = OxmlElement('w:fldChar')
            fldChar.set(qn('w:fldCharType'), 'begin')
            r.append(fldChar)
            instrText = OxmlElement('w:instrText')
            instrText.text = ' SEQ Figure \\* ARABIC'
            r.append(instrText)
            fldChar = OxmlElement('w:fldChar')
            fldChar.set(qn('w:fldCharType'), 'end')
            r.append(fldChar)
        
            paragraph.add_run(f": {text}")
            
        
        def add_equation_caption(location, text, doc):
        
            paragraph = doc.paragraphs[-1]
        
            drapeau = False
            for para in doc.paragraphs:
                if not drapeau:
                    if location in para.text:
                        paragraph = insert_paragraph_after(para, "Equation ", style="Caption")
                        drapeau = True
        
            run = paragraph.add_run()
            r = run._r
            fldChar = OxmlElement('w:fldChar')
            fldChar.set(qn('w:fldCharType'), 'begin')
            r.append(fldChar)
            instrText = OxmlElement('w:instrText')
            instrText.text = ' SEQ Figure \\* ARABIC'
            r.append(instrText)
            fldChar = OxmlElement('w:fldChar')
            fldChar.set(qn('w:fldCharType'), 'end')
            r.append(fldChar)
        
            paragraph.add_run(f": {text}")
            
        
        def remove_key_table(key, table):
            for row in table.rows:
                for cell in row.cells:
                    if key in cell.text:
                        cell.text = cell.text.replace('Key = ', '')
                        
        def remove_key_paragraph(key, pararaph):
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace('Key = ', '')

            
        def get_y_pred():
            y_pred = []
            for i in range (len(projects)):
                new_y = [projects[i]['coefficients']['intercept']]*len(projects[i]['Target']['normalized target'])            
                for n in range (len(projects[i]['Target']['normalized target'])):
                    for m in range (len(projects[i]['combination']['names'])):
                    
                        variable = projects[i]['combination']['names'][m]
                        new_y[n] += projects[i]['coefficients']['slopes'][m] * projects[i]['combination'][variable][n]
                    
                    new_y[n] = new_y[n]*projects[i]['Target']['timedelta'][n]
                y_pred.append(new_y)
            
            return y_pred
        
        y_pred = get_y_pred()
        
        
        # Table : Project details
        
        table = find_table('Key = Project Details')
        
        for row in table.rows:
        
            if 'Key = Project Details' in row.cells[1].text:
                row.cells[1].text = projects[0]['Scope']
            
            if 'ddress' in row.cells[0].text and projects[0]['Address']:
                row.cells[1].text = projects[0]['Address']
                                
            if 'M&V option' in row.cells[0].text and projects[0]['M&V option'] and projects[0]['Utility']['name']:
                option = projects[0]['M&V option'] + ' - '
                for i in range (len(projects)):
                    option += projects[i]['Utility']['name']
                    if i < len(projects) - 1:
                        option += ', '
                row.cells[1].text = option
            
            if 'umber of buildings' in row.cells[0].text and projects[0]['Number of buildings']:
                row.cells[1].text = format_str(projects[0]['Number of buildings'])
            
            if 'umber of floors' in row.cells[0].text and projects[0]['Number of floors including ground']:
                row.cells[1].text = format_str(projects[0]['Number of floors including ground'])
            
            if 'lectricity tariff' in row.cells[0].text and projects[0]['Base tariff without VAT']:
                row.cells[1].text = format_str(projects[0]["Base tariff without VAT"]['value'])+ ' '+ str(projects[0]["Base tariff without VAT"]['unit'])
                
            if 'uilt-up area' in row.cells[0].text and projects[0]['Built-up area']['value']:
                row.cells[1].text = format_str(projects[0]["Built-up area"]["value"])
                    
            if 'onstruction area' in row.cells[0].text and projects[0]['Gross floor conditioned area']['value']:
                row.cells[1].text = format_str(projects[0]["Gross floor conditioned area"]["value"])
            
            if 'year' in row.cells[0].text and projects[0]['Construction year']:
                row.cells[1].text = format_str(projects[0]["Construction year"])                    
        
        
        
        
        # Table : Total Baseline and Savings Summary
        
        key = 'Key = Total Baseline and Savings Summary'
        table = find_table(key)
        
        energy_cost = 0
        savings_cost = 0
        
        for i in range (len(projects)):
            if i > 0:
                table.add_row()
                
            row = table.rows[-1]
            row.cells[0].text = projects[i]['Utility']['name']
            row.cells[1].text = format_str(sum(projects[i]['Baseline']['baseline'])) + ' ' + projects[i]['Utility']['unit']
            row.cells[2].text = format_str(projects[i]['Savings']['value']) + ' ' + projects[i]['Savings']['unit']
            row.cells[3].text = format_str(100 * projects[i]['Savings']['value'] / sum(projects[i]['Baseline']['baseline'])) + '%'
            energy_cost += sum(projects[i]['Baseline']['baseline']) * projects[i]['Base tariff without VAT']['value']
            savings_cost += projects[i]['Savings']['value'] * projects[i]['Base tariff without VAT']['value']
        
        table.add_row()
        row = table.rows[-1]
        row.cells[0].text = 'Total cost'
        row.cells[1].text = format_str(energy_cost) +' '+ projects[i]['Base tariff without VAT']['unit'].split('/')[0]
        row.cells[2].text = format_str(savings_cost) +' '+ projects[i]['Base tariff without VAT']['unit'].split('/')[0]
        row.cells[3].text = format_str(100*savings_cost/energy_cost) + '%'

        remove_key_table(key, table)
        
    
        
    
        #  Table : Utilities dates
    
        key = "Key = Utilities dates"
        table = find_table(key)
        
        for i in range (len(projects)):
            if i > 0:
                table.add_row()

            row = table.rows[-1]
            row.cells[0].text = projects[i]['Utility']['name']
            row.cells[1].text = projects[i]['Baseline from'][:10]
            row.cells[2].text = projects[i]['Baseline to'][:10]
            
        remove_key_table(key, table)
        
    
        
    
        # Table : All information for the baseline(s)
    
        original_table = find_table('Key = All baseline')
        
        location = "Key = Below are the tables corresponding to the baseline"
        
        for i in range (0, len(projects)):
            paragraph = find_paragraph(location)
            new_caption = 'Baseline for ' + projects[i]['Utility']['name']
            
            add_table_caption(paragraph,
                              new_caption,
                              doc)
            location = new_caption

        for i in range(0, len(projects)):
            
            caption = 'Baseline for ' + projects[i]['Utility']['name']
            paragraph = find_paragraph(caption)
            copy_table_after(original_table, paragraph)
            paragraph.insert_paragraph_before('')
            
            table = find_table('Key = All baseline')
            table.rows[0].cells[3].text = projects[i]['Utility']['name'] + ' (' + projects[i]['Utility']['unit'] + ')'
                
            for k in range(len(projects[i]['Baseline']['baseline'])):
                if k > 0:
                    table.add_row()
                    
                row = table.rows[-1]
                row.cells[0].text = str(k+1)
                row.cells[1].text = projects[i]['Baseline']['from'][k][:10]
                row.cells[2].text = projects[i]['Baseline']['to'][k][:10]
                row.cells[3].text = format_str(projects[i]['Baseline']['baseline'][k])
            
        original_table._element.getparent().remove(original_table._element)

        paragraph = find_paragraph(location)
        remove_key_paragraph(location, paragraph)




        # Table : Regression summary
        
        table = find_table('Key = regression summary')
        
        for i in range(len(projects)):
            if i>0:
                table.add_row()
                
            row = table.rows[-1]
            row.cells[0].text = projects[i]['Utility']['name']
            row.cells[2].text = format_str(len(projects[i]['Target']['normalized target']))
            
            text= ''
            for k in range(len(projects[i]['combination']['names'])):
                text += projects[i]['combination']['names'][k]
                if k < len(projects[i]['combination']['names'])-1:
                    text += ', '
            row.cells[1].text = text
        
        remove_key_table(key, table)            
        
        
        
        
        # Table : Data of the variables for each utility
        
        original_table = find_table('Key = variables data for each utility')
        key = "Key = The data for those variables is listed below along with the hourly-normalized baseline."
        location = key
        
        for i in range (0, len(projects)):
            paragraph = find_paragraph(location)
            new_caption = 'Variables data for ' + projects[i]['Utility']['name']
            
            add_table_caption(paragraph,
                              new_caption,
                              doc)
            location = new_caption


        for i in range(0, len(projects)):
             
            caption = 'Variables data for ' + projects[i]['Utility']['name']
            paragraph = find_paragraph(caption)
            copy_table_after(original_table, paragraph)
            if i>0:
                paragraph.insert_paragraph_before('')
            
            table = find_table('Key = variables data for each utility')
            
            row = table.rows[0]
            row.cells[3].text = projects[i]['Utility']['name'] + ' (' + projects[i]['Utility']['unit'] + ')'
            
            nb_of_columns_to_delete = 4-len(projects[i]['combination']['names'])
            
            for j in range (nb_of_columns_to_delete):
                delete_column_in_table(table, 4)
                  
            for j in range (len(projects[i]['combination']['names'])):
                row.cells[j+4].text = projects[i]['combination']['names'][j]

            for k in range(len(projects[i]['Target']['target'])):
                if k > 0:
                    table.add_row()

                row = table.rows[-1]
                row.cells[0].text = str(k+1)
                row.cells[1].text = projects[i]['Target']['from'][k][:10]
                row.cells[2].text = projects[i]['Target']['to'][k][:10]
                row.cells[3].text = format_str(projects[i]['Target']['normalized target'][k])
                
                for j in range (len(projects[i]['combination']['names'])):
                    variable = projects[i]['combination']['names'][j]
                    row.cells[j+4].text = format_str(projects[i]['combination'][variable][k])
                    
        original_table._element.getparent().remove(original_table._element)
        paragraph = find_paragraph(key)
        remove_key_paragraph(key, paragraph)            
        
        
        
        
        # Add number of total results to pararaph
    
        total_nb_of_results = sum([projects[i]['number of results'] for i in range (len(projects))])
        key = 'Key = number of results'
        paragraph = find_paragraph(key)
        paragraph.text = paragraph.text.replace(key, format_str(total_nb_of_results))
    
    
    
    
        # Regression equation 
    
        key = 'Key = The regression equation of the different models are'
        paragraph = find_paragraph(key)

        for i in range (len(projects)):
            equation = projects[i]['equation']
            utility = projects[i]['Utility']['name']
            unit = projects[i]['Utility']['unit']
            
            paragraph = insert_paragraph_after(paragraph, "For "+ utility)
            equation = equation.replace('normalized baseline', 'baseline (' + unit + ")")
            equation = equation.replace('= ', '= (')
            equation += ') * nb_of_hours'
            
            EXPRESSION = r"{}".format(equation)
        
            parser=mathtext.MathTextParser( 'bitmap' )
            offset=parser.to_png("equation.png", EXPRESSION, fontsize=12)
            
            r = paragraph.add_run("")
            inline_obj= r.add_picture("equation.png", width = Inches(7.0), height = Inches(0.2)) #return Inline object
            
        paragraph = find_paragraph(key)
        remove_key_paragraph(key, paragraph)
        
        


        #  Scatter plots
        
        key = "Key = The scatter plots below"
        location = key
        
        for i in range (0, len(projects)):
            paragraph = find_paragraph(location)
            insert_paragraph_after(paragraph, "Figure"+str(i))
            location = "Figure"+str(i) 


        for i in range (0, len(projects)):

            y_pred_i = y_pred[i]
            fig = plt.figure(figsize = (7, 5))
            plt.scatter(y_pred_i, projects[i]['Target']['target'], color = 'blue')
            plt.plot(projects[i]['Target']['target'], projects[i]['Target']['target'], color = 'red')
            plt.ylabel('Predictions')
            plt.xlabel('baseline (' + projects[i]['Utility']['unit'] + ')')
            plt.title('Predictions as a function of the baseline for ' + projects[i]['Utility']['name'])
            
            plt.savefig(str(i)+'.png')
        
            paragraph = find_paragraph("Figure"+str(i))
            
            new_caption = 'Predictions as a function of the baseline for ' + projects[i]['Utility']['name']
            add_figure_caption("Figure"+str(i),
                               new_caption,
                               doc)
            paragraph = find_paragraph("Figure"+str(i))
            paragraph.text = paragraph.text.replace("Figure"+str(i), "")
            run = paragraph.add_run()
            run.add_picture(str(i)+'.png', width = Inches(7.0), height = Inches(4.0))
            os.remove(str(i)+'.png')
        
        paragraph = find_paragraph(key)
        remove_key_paragraph(key, paragraph)
        
        
        
        
        # Table : IPMVP criteria

        original_table = find_table('Key = IPMVP criteria')
        
        key = "Key = IPMVP defines the below"
        location = key
        
        for i in range (0, len(projects)):
            paragraph = find_paragraph(location)
            new_caption = 'IPMVP criteria for ' + projects[i]['Utility']['name']
            
            add_table_caption(paragraph,
                              new_caption,
                              doc)
            location = new_caption


        for i in range(0, len(projects)):
             
            caption = 'IPMVP criteria for ' + projects[i]['Utility']['name']
            paragraph = find_paragraph(caption)
            copy_table_after(original_table, paragraph)
            if i>0:
                paragraph.insert_paragraph_before('')
            
            table = find_table('Key = IPMVP criteria')
            
            row0 = table.rows[0]
            row0.cells[0].text = projects[i]['Utility']['name']
            
            row1 = table.rows[1]
            row1.cells[1].text = format_str(projects[i]['r2'])
            if projects[i]['r2'] > 0.75:
                row1.cells[-1].text = 'yes'
            else:
                row1.cells[-1].text = 'no'
            
            row2 = table.rows[2]
            row2.cells[1].text = format_str(projects[i]['cv_rmse'])
            if projects[i]['cv_rmse'] < 0.2:
                row2.cells[-1].text = 'yes'
            else:
                row2.cells[-1].text = 'no'
            
            row3 = table.rows[3]
            row3.cells[1].text = format_str(projects[i]['tval'][0])
            if abs(projects[i]['tval'][0]) > 2:
                row3.cells[-1].text = 'yes'
            else:
                row3.cells[-1].text = 'no'
            
            for k in range(len(projects[i]['combination']['names'])):
                table.add_row()
                row = table.rows[-1]
                row.cells[0].text = 't-stat, ' + projects[i]['combination']['names'][k]
                row.cells[1].text = format_str(projects[i]['tval'][k+1])
                row.cells[2].text = row3.cells[2].text
                if abs(projects[i]['tval'][k+1]) > 2:
                    row.cells[-1].text = 'yes'
                else:
                    row.cells[-1].text = 'no'
                    
        original_table._element.getparent().remove(original_table._element)
        paragraph = find_paragraph(key)
        remove_key_paragraph(key, paragraph)
        
        
        
        
        # Table : Standard error
        
        key = 'Key = standard error'
        table = find_table(key)
        uncertainty = []
        std_total = []
        
        for i in range(len(projects)):
            if i>0:
                table.add_row()
            
            row = table.rows[-1]
            
            std_total.append(projects[i]['std_dev']*np.sqrt(len(projects[i]['Target']['normalized target'])))
            uncertainty.append(100*std_total[i]/projects[i]['Savings']['value'])
            
            row.cells[0].text = projects[i]['Scope']
            row.cells[1].text = projects[i]['Utility']['name']
            row.cells[2].text = format_str(projects[i]['std_dev'])
            row.cells[3].text = format_str(std_total[i])
            row.cells[4].text = format_str(projects[i]['Savings']['value'])
            row.cells[5].text = format_str(uncertainty[i]) + '%'
        
        remove_key_table(key, table)
        
        
        
        
        # Table : Confidence and precision
        
        key = 'Key = confidence'
        table = find_table(key)
        
        plus_or_minus = table.rows[1].cells[1].text
        
        for i in range(len(projects)):
            if i>0:
                table.add_row()
            
            row = table.rows[-1]
            
            z = scipy.stats.t.ppf(0.95, len(projects[i]['Target']['normalized target']) - len(projects[i]['combination']['names']) - 1) 
            relative_precision = z * uncertainty[i]
            absolute_precision = z * std_total[i]
            
            row.cells[0].text = projects[i]['Scope']
            row.cells[1].text = projects[i]['Utility']['name']
            row.cells[2].text = plus_or_minus + format_str(relative_precision)
            row.cells[3].text = format_str(absolute_precision)
            row.cells[4].text = format_str(projects[i]['Savings']['value'] - absolute_precision)
            row.cells[5].text = format_str(projects[i]['Savings']['value'] + absolute_precision)
            row.cells[6].text = format_str(z)
        
        remove_key_table(key, table)

        


        # Characteristics graphs
        
        key = 'Key = Below are the baseline characteristics graphs'
        location = key
        
        for i in range (0, len(projects)):
            for k in range (0, len(projects[i]['combination']['names'])):
                paragraph = find_paragraph(location)
                insert_paragraph_after(paragraph, "Figure"+str(i)+str(k))
                location = "Figure"+str(i)+str(k)

        
        for i in range (0, len(projects)):
            
            for k in range (0, len(projects[i]['combination']['names'])):
                
                trace1 = go.Bar(
                    x= projects[i]['Target']['from'],
                    y = projects[i]['Target']['normalized target'],
                    name = 'Baseline ('+ projects[i]['Utility']['unit'] + ')',
                    )
                
                variable = projects[i]['combination']['names'][k]
                
                trace2 = go.Scatter(
                    x = projects[i]['Target']['from'],
                    y = projects[i]['combination'][variable],
                    name = variable
                    )
                
                fig = make_subplots(specs=[[{"secondary_y": True}]])
                fig.add_trace(trace1)
                fig.add_trace(trace2,secondary_y=True)
                fig['layout'].update(xaxis=dict(
                      tickangle=-90
                    ))
                fig.update_layout(
                    title = {'text' : 'Baseline and '+variable+' for '+ projects[i]['Utility']['name'],'x':0.47, 'xanchor': 'center', 'yanchor': 'top'}, 
                    xaxis_title ='Time', 
                    yaxis_title='Baseline ('+ projects[i]['Utility']['unit'] + ')',
                    legend=dict(yanchor="top", y=0.99, xanchor="left", x=0.01)
                    )
                fig.update_yaxes(title_text=units[variable], secondary_y=True)
                fig.write_image(str(i)+str(k)+'.png')
                
                paragraph = find_paragraph("Figure"+str(i)+str(k))
                
                variable_for_caption = projects[i]['combination']['names'][k]
                
                new_caption = 'Baseline and '+ variable_for_caption  +' for '+ projects[i]['Utility']['name']
                add_figure_caption("Figure"+str(i)+str(k),
                                   new_caption,
                                   doc)
    
                paragraph = find_paragraph("Figure"+str(i)+str(k))
                paragraph.text = paragraph.text.replace("Figure"+str(i)+str(k), "")
                run = paragraph.add_run()
                run.add_picture(str(i)+str(k)+'.png', width = Inches(7.0), height = Inches(4.0))
                os.remove(str(i)+str(k)+'.png')

        paragraph = find_paragraph(key)
        remove_key_paragraph(key, paragraph)




        # Change the tables style to 'No Spacing'

        styles = doc.styles
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:                
                    for paragraph in cell.paragraphs:
                        paragraph.style = styles['No Spacing']


        # Save the file
        doc.save('Final_Report.docx')
        
        from io import BytesIO
        
        f = BytesIO()
        doc.save(f)
        st.download_button(label = "Download Report",
                           data = f,
                           file_name = "Final_Report.docx")
        
        # st.download_button(label="Download PDF Tutorial", 
        #         data=PDFbyte,
        #         file_name="Final_report.pdf",
        #         mime='application/octet-stream')
        
        
        # doc.save(r'C:\Users\Bruno Tabet\Documents\ENOVA\MVP\M&VTemplate.docx')
        
        
        # def get_repo_name_from_url(url: str) -> str:
        #     last_slash_index = url.rfind("/")
        #     last_suffix_index = url.rfind(".git")
        #     if last_suffix_index < 0:
        #         last_suffix_index = len(url)
        
        #     if last_slash_index < 0 or last_suffix_index <= last_slash_index:
        #         raise Exception("Badly formatted url {}".format(url))
        
        #     return url[last_slash_index + 1:last_suffix_index]
                
        # PATH_OF_GIT_REPO = get_repo_name_from_url('https://github.com/BrunoTabet/MVPPublic5')  # make sure .git folder is properly configured
        # COMMIT_MESSAGE = 'comment from python script'
        
        # def git_push():
        #     try:
        #         repo = Repo(str(PATH_OF_GIT_REPO))
        #         repo.git.add(update=True)
        #         repo.index.commit(COMMIT_MESSAGE)
        #         origin = repo.remote(name='origin')
        #         origin.push()
        #     except:
        #         st.write('Some error occured while pushing the code')
        #         # st.stop()
        
        # git_push()
        
        
        # doc.save('Final_report.pdf')
        
        # def show_pdf(file_path):
        #     with open(file_path,"rb") as f:
        #         base64_pdf = base64.b64encode(f.read()).decode('utf-8')
        #     pdf_display = f'<iframe src="data:application/pdf;base64,{base64_pdf}" width="800" height="800" type="application/pdf"></iframe>'
        #     st.markdown(pdf_display, unsafe_allow_html=True)
        
        # if st.button('PDFFFF'):
        #     show_pdf('Final_report.pdf')
        
        # # with open("Final_report.pdf", "rb") as pdf_file:
        # #     # st.markdown(pdf_file, unsafe_allow_html=True)
        # #     PDFbyte = pdf_file.read()
        # #     # st.markdown(PDFbyte, unsafe_allow_html =True)
        
        # # st.download_button(label="Download PDF Tutorial", 
        # #         data=PDFbyte,
        # #         file_name="Final_report.pdf",
        # #         mime='application/octet-stream')
        
        
        # with open("Final_report.pdf", "rb") as file:
        #     btn=st.download_button(
        #     label="click me to download pdf",
        #     data=file,
        #     file_name="ok.pdf",
        #     mime="application/octet-stream"
        # )
        
        
        # # from fpdf import FPDF
        # # import base64
        
        # # pdf = FPDF() #pdf object
        # # pdf=FPDF(orientation="P", unit="mm", format="A4")
        # # pdf.add_page()
        
        # # pdf.set_font("Times", "B", 18)
        # # pdf.set_xy(10.0, 20)
    
        # # pdf.cell(w = 75.0, h = 5.0, align = "L", txt = "This is my sample text")
        
        # # oflnme = "Final_report.pdf"
        # # b64 = base64.b64encode(pdf.output(dest='S')).decode()
        # # st.download_button("Download Report", data=b64, file_name=oflnme, mime="application/octet-stream", help=f"Download file {oflnme}")
        
        # #alternatively, if I replace the above last 2 lines with the following 3 lines of code for a download hyperlink, it works fine.
        
        # # b64 = base64.b64encode(pdf.output(dest="S"))
        # # html = f'Download file'
        # # st.markdown(html, unsafe_allow_html=True)



        # import base64
        # import os
        # import json
        # import pickle
        # import uuid
        # import re
        
        # import streamlit as st
        # import pandas as pd
        
        
        # def download_button(object_to_download, download_filename, button_text, pickle_it=False):
        #     """
        #     Generates a link to download the given object_to_download.
        
        #     Params:
        #     ------
        #     object_to_download:  The object to be downloaded.
        #     download_filename (str): filename and extension of file. e.g. mydata.csv,
        #     some_txt_output.txt download_link_text (str): Text to display for download
        #     link.
        #     button_text (str): Text to display on download button (e.g. 'click here to download file')
        #     pickle_it (bool): If True, pickle file.
        
        #     Returns:
        #     -------
        #     (str): the anchor tag to download object_to_download
        
        #     Examples:
        #     --------
        #     download_link(your_df, 'YOUR_DF.csv', 'Click to download data!')
        #     download_link(your_str, 'YOUR_STRING.txt', 'Click to download text!')
        
        #     """
        #     if pickle_it:
        #         try:
        #             object_to_download = pickle.dumps(object_to_download)
        #         except pickle.PicklingError as e:
        #             st.write(e)
        #             return None
        
        #     else:
        #         if isinstance(object_to_download, bytes):
        #             pass
        
        #         elif isinstance(object_to_download, pd.DataFrame):
        #             object_to_download = object_to_download.to_csv(index=False)
        
        #         # Try JSON encode for everything else
        #         else:
        #             object_to_download = json.dumps(object_to_download)
        
        #     try:
        #         # some strings <-> bytes conversions necessary here
        #         b64 = base64.b64encode(object_to_download.encode()).decode()
        
        #     except AttributeError as e:
        #         b64 = base64.b64encode(object_to_download).decode()
        
        #     button_uuid = str(uuid.uuid4()).replace('-', '')
        #     button_id = re.sub('\d+', '', button_uuid)
        
        #     custom_css = f""" 
        #         <style>
        #             #{button_id} {{
        #                 background-color: rgb(255, 255, 255);
        #                 color: rgb(38, 39, 48);
        #                 padding: 0.25em 0.38em;
        #                 position: relative;
        #                 text-decoration: none;
        #                 border-radius: 4px;
        #                 border-width: 1px;
        #                 border-style: solid;
        #                 border-color: rgb(230, 234, 241);
        #                 border-image: initial;
        
        #             }} 
        #             #{button_id}:hover {{
        #                 border-color: rgb(246, 51, 102);
        #                 color: rgb(246, 51, 102);
        #             }}
        #             #{button_id}:active {{
        #                 box-shadow: none;
        #                 background-color: rgb(246, 51, 102);
        #                 color: white;
        #                 }}
        #         </style> """
        
        #     dl_link = custom_css + f'<a download="{download_filename}" id="{button_id}" href="data:file/txt;base64,{b64}">{button_text}</a><br></br>'
        
        #     return dl_link
        
        
        # from email.message import EmailMessage

        
        # # Document = document()
        # # paragraph = document.add_paragraph("Test Content")
        # f = BytesIO()
        # doc.save(f)
        # file_list = []
        # file_list.append(["Final_Report.docx", f.get_value(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"])
        # email = EmailMessage(subject = 'Test', body = 'Hi', to = ['barbourjapon@gmail.com'], attachments = file_list)
        # email.send()
        
        
        
        # # download_button(doc, 'Final_Report.docx', 'DOWNLOAD HERE')

        
        st.session_state['M&V'] = 2
        

        
if st.session_state['M&V'] == 2:
    st.write('Your report has been created successfully !')
    
    # link = 'https://github.com/BrunoTabet/MVPPublic5/blob/master/M&VReport.docx'
    link = 'https://github.com/BrunoTabet/MVPPublic5/raw/master/M%VReport.docx'
    link = 'https://github.com/BrunoTabet/MVPPublic5/raw/master/Final_Report.docx'
    path = r"C:\Users\Bruno Tabet\Documents\ENOVA\MVP\THISISATEST.docx"
    # path = 'M&VReport.docx'
    st.write("You can download the report by clicking [here](%s)." % link)
    
    st.write('')
    st.write('')
    st.write('')
    
    # with open("M&VReport.pdf", "rb") as pdf_file:
    #     PDFbyte = pdf_file.read()
    
    # st.download_button(label = 'OK', data = pdf_file, file_name = 'M&VReport.pdf')
    
    
    if st.button('Create another report', key = 675675967567):
        st.session_state['M&V'] = 1
        st.experimental_rerun()
        
    st.write('')
    
    col1, col2, col3 = st.columns([1, 5, 1])        

    with col1:
        if st.button("< Prev"):
            nav_page('Database')